"""
ingestion/document_loader.py
────────────────────────────
Responsible for loading raw text out of PDF and DOCX files.

Supports two PDF backends:
  • PyMuPDF  (fitz) – fast, accurate, handles scanned PDFs via OCR if needed.
  • pdfplumber          – better for tables and complex layouts.

The caller always receives a plain Python string regardless of the file type.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)


# ─── PDF helpers ──────────────────────────────────────────────────────────────

def _load_pdf_pymupdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).
    This is the primary PDF parser – fast and reliable.
    """
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    # Open from bytes so we never need a disk path
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")  # type: ignore[attr-defined]
            if page_text.strip():
                # Prepend a page marker so citations can reference it later
                text_parts.append(f"[Page {page_num}]\n{page_text}")

    return "\n\n".join(text_parts)


def _load_pdf_pdfplumber(file_bytes: bytes) -> str:
    """
    Fallback PDF parser using pdfplumber.
    Better at extracting text from tables and multi-column layouts.
    """
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")

    return "\n\n".join(text_parts)


# ─── DOCX helper ──────────────────────────────────────────────────────────────

def _load_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Iterates over paragraphs and table cells to capture all text.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Main body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables (contracts often have tables with key terms)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


# ─── Public interface ─────────────────────────────────────────────────────────

def load_document(
    file_bytes: bytes,
    filename: str,
) -> str:
    """
    Load a document from raw bytes and return its full text content.

    Parameters
    ----------
    file_bytes : bytes
        The raw binary content of the uploaded file.
    filename : str
        Original filename – used only to determine the file type.

    Returns
    -------
    str
        Extracted text, with page markers where applicable.

    Raises
    ------
    ValueError
        If the file extension is not supported.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        logger.info("Loading PDF with PyMuPDF: %s", filename)
        try:
            text = _load_pdf_pymupdf(file_bytes)
            if not text.strip():
                # If PyMuPDF returned nothing, fall back to pdfplumber
                logger.warning("PyMuPDF returned empty text; trying pdfplumber fallback.")
                text = _load_pdf_pdfplumber(file_bytes)
            return text
        except Exception as exc:
            logger.warning("PyMuPDF failed (%s); trying pdfplumber.", exc)
            return _load_pdf_pdfplumber(file_bytes)

    elif suffix == ".docx":
        logger.info("Loading DOCX: %s", filename)
        return _load_docx(file_bytes)

    else:
        raise ValueError(
            f"Unsupported file type: '{suffix}'. "
            "Please upload a .pdf or .docx file."
        )


def load_document_from_path(file_path: Union[str, Path]) -> str:
    """
    Convenience wrapper: load a document by its file-system path.
    Useful for CLI tests and the ingestion pipeline.
    """
    path = Path(file_path)
    file_bytes = path.read_bytes()
    return load_document(file_bytes, path.name)
